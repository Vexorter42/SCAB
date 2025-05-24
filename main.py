from fastapi import FastAPI
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from docx import Document
import random
import os

app = FastAPI()

@app.get("/", response_class=HTMLResponse)
def read_root():
    with open("app/templates/index.html", "r", encoding="utf-8") as f:
        return f.read()

def generate_resume_data():
    first_names = ["Алексей", "Михаил", "Иван", "Олег", "Дмитрий", "Евдоким"]
    last_names = ["Иванов", "Петров", "Сидоров", "Смирнов", "Кузнецов", "Попов"]
    jobs = ["Разработчик", "Аналитик", "Менеджер проектов", "Тестировщик", "UX/UI дизайнер"]
    skills = ["Python", "JavaScript", "SQL", "Django", "React", "Git", "Docker", "Agile"]

    name = f"{random.choice(first_names)} {random.choice(last_names)}"
    job = random.choice(jobs)
    experience = random.randint(1, 10)
    selected_skills = random.sample(skills, 4)

    return {
        "name": name,
        "job": job,
        "experience": experience,
        "skills": selected_skills
    }

def generate_markdown(data):
    return f"""# Резюме

**Имя:** {data['name']}

**Профессия:** {data['job']}

**Опыт:** {data['experience']} лет

**Навыки:** {", ".join(data['skills'])}
"""

def generate_docx(data, filename):
    doc = Document()
    doc.add_heading('Резюме', 0)
    doc.add_paragraph(f"Имя: {data['name']}")
    doc.add_paragraph(f"Профессия: {data['job']}")
    doc.add_paragraph(f"Опыт: {data['experience']} лет")
    doc.add_paragraph(f"Навыки: {', '.join(data['skills'])}")
    doc.save(filename)

@app.get("/resume/markdown")
def get_markdown_resume():
    data = generate_resume_data()
    md = generate_markdown(data)
    with open("resume.md", "w", encoding="utf-8") as f:
        f.write(md)
    return FileResponse("resume.md", media_type="text/markdown", filename="resume.md")

@app.get("/resume/docx")
def get_docx_resume():
    data = generate_resume_data()
    generate_docx(data, "resume.docx")
    return FileResponse("resume.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="resume.docx")
